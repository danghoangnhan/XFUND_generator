"""
DOCX template processing utilities.
Handles filling DOCX templates with CSV data and converting to PNG images.
"""

import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING, Optional

from docx import Document
from pdf2image import convert_from_path
from PIL import Image

from .models import TemplateValidationResult
from .utils import ensure_dir_exists, normalize_field_name

if TYPE_CHECKING:
    from docx.document import Document as DocumentType

logger = logging.getLogger(__name__)

# Placeholder patterns for template field detection
PLACEHOLDER_PATTERNS = [
    r"\{\{([^}]+)\}\}",  # {{field}}
    r"\{([^}]+)\}",  # {field}
    r"\[([^\]]+)\]",  # [field]
]


class DocxProcessor:
    """Handles DOCX template processing and conversion."""

    def __init__(self, template_path: str):
        """
        Initialize DOCX processor with template.

        Args:
            template_path: Path to the DOCX template file
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        self.template_path = template_path
        self.template_name = Path(template_path).stem

    def fill_template(
        self, data: dict[str, str], output_path: Optional[str] = None
    ) -> str:
        """
        Fill DOCX template with data using placeholder replacement.

        Args:
            data: Dictionary containing field values
            output_path: Optional path for output DOCX file

        Returns:
            Path to the filled DOCX file
        """
        if output_path is None:
            output_path = os.path.join(
                tempfile.gettempdir(), f"filled_{self.template_name}_{id(data)}.docx"
            )

        # Load template document
        doc = Document(self.template_path)

        # Replace placeholders in paragraphs
        self._replace_placeholders_in_paragraphs(doc, data)

        # Replace placeholders in tables
        self._replace_placeholders_in_tables(doc, data)

        # Replace placeholders in headers/footers
        self._replace_placeholders_in_headers_footers(doc, data)

        # Save filled document
        ensure_dir_exists(os.path.dirname(output_path))
        doc.save(output_path)

        logger.info(f"Filled template saved to: {output_path}")
        return output_path

    def _replace_placeholders_in_paragraphs(
        self, doc: "DocumentType", data: dict[str, str]
    ) -> None:
        """Replace placeholders in document paragraphs."""
        for paragraph in doc.paragraphs:
            for field_name, value in data.items():
                placeholder = self._create_placeholder(field_name)
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _replace_placeholders_in_tables(
        self, doc: "DocumentType", data: dict[str, str]
    ) -> None:
        """Replace placeholders in document tables."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field_name, value in data.items():
                        placeholder = self._create_placeholder(field_name)
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

    def _replace_placeholders_in_headers_footers(
        self, doc: "DocumentType", data: dict[str, str]
    ) -> None:
        """Replace placeholders in headers and footers."""
        for section in doc.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                for field_name, value in data.items():
                    placeholder = self._create_placeholder(field_name)
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                for field_name, value in data.items():
                    placeholder = self._create_placeholder(field_name)
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _create_placeholder(self, field_name: str) -> str:
        """
        Create placeholder string for field name.
        Supports multiple placeholder formats: {{field}}, {field}, [field]
        """
        # Try normalized field name first
        normalized_name = normalize_field_name(field_name)
        placeholders = [
            f"{{{{{normalized_name}}}}}",  # {{field}}
            f"{{{normalized_name}}}",  # {field}
            f"[{normalized_name}]",  # [field]
            f"{{{{{field_name}}}}}",  # {{original}}
            f"{{{field_name}}}",  # {original}
            f"[{field_name}]",  # [original]
        ]
        return placeholders[0]  # Default to {{field}} format

    def convert_to_pdf(self, docx_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert DOCX to PDF using LibreOffice.

        Args:
            docx_path: Path to the DOCX file
            output_path: Optional path for output PDF file

        Returns:
            Path to the generated PDF file
        """
        if output_path is None:
            output_path = docx_path.replace(".docx", ".pdf")

        output_dir = os.path.dirname(output_path)
        ensure_dir_exists(output_dir)

        try:
            # Use LibreOffice to convert DOCX to PDF
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path,
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

            # LibreOffice saves PDF with same basename as input
            expected_pdf = os.path.join(
                output_dir, os.path.basename(docx_path).replace(".docx", ".pdf")
            )

            if expected_pdf != output_path and os.path.exists(expected_pdf):
                shutil.move(expected_pdf, output_path)

            if not os.path.exists(output_path):
                raise RuntimeError(f"PDF file not created: {output_path}")

            logger.info(f"PDF created: {output_path}")
            return output_path

        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out") from None
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            raise

    def convert_to_png(
        self, docx_path: str, output_path: Optional[str] = None, dpi: int = 300
    ) -> str:
        """
        Convert DOCX to PNG via PDF conversion.

        Args:
            docx_path: Path to the DOCX file
            output_path: Optional path for output PNG file
            dpi: DPI for PNG conversion (higher = better quality)

        Returns:
            Path to the generated PNG file
        """
        if output_path is None:
            output_path = docx_path.replace(".docx", ".png")

        # First convert to PDF
        pdf_path = docx_path.replace(".docx", ".pdf")
        try:
            pdf_path = self.convert_to_pdf(docx_path, pdf_path)

            # Then convert PDF to PNG
            png_path = self.pdf_to_png(pdf_path, output_path, dpi)

            # Clean up temporary PDF
            if os.path.exists(pdf_path) and pdf_path != png_path.replace(
                ".png", ".pdf"
            ):
                os.remove(pdf_path)

            return png_path

        except Exception as e:
            logger.error(f"Error converting DOCX to PNG: {e}")
            raise

    def pdf_to_png(self, pdf_path: str, output_path: str, dpi: int = 300) -> str:
        """
        Convert PDF to PNG using pdf2image.

        Args:
            pdf_path: Path to the PDF file
            output_path: Path for output PNG file
            dpi: DPI for conversion

        Returns:
            Path to the generated PNG file
        """
        try:
            # Convert first page of PDF to image
            images = convert_from_path(pdf_path, dpi=dpi, first_page=1, last_page=1)

            if not images:
                raise RuntimeError(f"No images generated from PDF: {pdf_path}")

            # Save as PNG
            ensure_dir_exists(os.path.dirname(output_path))
            images[0].save(output_path, "PNG")

            logger.info(f"PNG created: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error converting PDF to PNG: {e}")
            raise

    def get_image_size(self, image_path: str) -> tuple[int, int]:
        """
        Get image dimensions.

        Args:
            image_path: Path to the image file

        Returns:
            Tuple of (width, height)
        """
        try:
            with Image.open(image_path) as img:
                return img.size
        except Exception as e:
            logger.error(f"Error getting image size: {e}")
            raise


def process_docx_template(
    template_path: str,
    data: dict[str, str],
    output_image_path: str,
    dpi: int = 300,
    cleanup_temp: bool = True,
) -> tuple[str, tuple[int, int]]:
    """
    Complete pipeline: fill DOCX template and convert to PNG.

    Args:
        template_path: Path to DOCX template
        data: Data to fill in template
        output_image_path: Path for output PNG
        dpi: DPI for image conversion
        cleanup_temp: Whether to clean up temporary files

    Returns:
        Tuple of (image_path, image_size)
    """
    processor = DocxProcessor(template_path)

    # Create temporary filled DOCX
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_docx = os.path.join(temp_dir, f"filled_{processor.template_name}.docx")

        # Fill template
        filled_docx = processor.fill_template(data, temp_docx)

        # Convert to PNG
        image_path = processor.convert_to_png(filled_docx, output_image_path, dpi)

        # Get image size
        image_size = processor.get_image_size(image_path)

        return image_path, image_size


def check_libreoffice_installed() -> bool:
    """
    Check if LibreOffice is installed and available.

    Returns:
        True if LibreOffice is available, False otherwise
    """
    try:
        result = subprocess.run(
            ["libreoffice", "--version"], capture_output=True, text=True, timeout=10
        )
        return result.returncode == 0
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False


def validate_docx_template(template_path: str) -> TemplateValidationResult:
    """
    Validate DOCX template and extract information.

    Args:
        template_path: Path to DOCX template

    Returns:
        TemplateValidationResult with template validation results
    """
    if not os.path.exists(template_path):
        return TemplateValidationResult.create_error("File not found")

    try:
        doc = Document(template_path)

        # Extract all text to find placeholders
        all_text = []
        for paragraph in doc.paragraphs:
            all_text.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        full_text = " ".join(all_text)

        # Find potential placeholders using module-level patterns
        placeholders = []
        for pattern in PLACEHOLDER_PATTERNS:
            matches = re.findall(pattern, full_text)
            placeholders.extend(matches)

        return TemplateValidationResult.create_success(
            placeholders=list(set(placeholders)),
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables),
        )

    except Exception as e:
        return TemplateValidationResult.create_error(str(e))
